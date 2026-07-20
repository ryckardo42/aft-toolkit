#!/usr/bin/env python3
"""gera_relatorio_docx.py — Relatório Final Simplificado (.docx) do /sfitweb-rel.

Uso:
    python3 gera_relatorio_docx.py <relatorio-final.json> [saida.docx]

Lê os dados estruturados do relatório e gera o .docx sobre o template oficial
com o cabeçalho da auditoria (scripts/template-cabecalho.docx — cópia de
Template/Template com cabeçalho.docx do toolkit), aplicando o modelo
"Relatório de Fiscalização / Indícios": Times New Roman 12, margens
2/2/2/2,5 cm, paleta azul (#1F3864 / #2E5496), corpo justificado 1,15,
tabelas com cabeçalho azul, subcabeçalho por tema e zebra #EBF3FB/#F5F5F5.

Esquema do JSON (todas as chaves de texto são strings simples; listas vazias
ou chaves ausentes fazem a seção correspondente ser omitida/ajustada):
{
  "titulo":          "RELATÓRIO FINAL SIMPLIFICADO",
  "subtitulo":       "Relatório de fiscalização trabalhista — consolidação da ação fiscal",
  "unidade":         "EMPRESA LTDA — CNPJ 00.000.000/0000-00",
  "data":            "Goiânia-GO, 20 de julho de 2026",
  "identificacao":   [["Empresa Fiscalizada", "..."], ["CNPJ", "..."], ...],
  "sintese":         ["parágrafo 1", "parágrafo 2"],
  "notificacoes":    [{"codigo": "...", "tipo": "NAD", "itens": "resumo...", "lavrada": "dd/mm/aaaa"}],
  "autos_temas":     [{"tema": "...", "autos": [{"numero": "...", "ementa": "...",
                       "fundamento": "...", "descricao": "...", "constatacao": "..."}]}],
  "autos_total":     "Total: 20 autos de infração transmitidos.",
  "sem_autos":       "frase única quando não há autos transmitidos (autos_temas vazio)",
  "interdicoes":     ["parágrafo 1", ...]  (vazio → frase padrão de inexistência),
  "observacoes":     ["item 1", "item 2"]  (vazio → "Nenhuma pendência identificada."),
  "auditores":       [["Ricardo de Oliveira", "Auditor-Fiscal do Trabalho — CIF 000000"]]
}
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AZUL_ESCURO = "1F3864"
AZUL_MEDIO = "2E5496"
BRANCO = "FFFFFF"
CINZA_CAPA1 = "444444"
CINZA_CAPA2 = "555555"
ZEBRA_AZUL = "EBF3FB"
ZEBRA_CINZA = "F5F5F5"
BORDA = "AAAAAA"

FONTE = "Times New Roman"
TAM = Pt(12)


def rgb(hexa: str) -> RGBColor:
    return RGBColor.from_string(hexa)


def _fmt(p, *, antes=0, depois=10, alinh=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(antes)
    pf.space_after = Pt(depois)
    pf.alignment = alinh
    return p


def _run(p, texto, *, negrito=False, italico=False, cor="000000"):
    r = p.add_run(texto)
    r.font.name = FONTE
    r.font.size = TAM
    r.bold = negrito
    r.italic = italico
    r.font.color.rgb = rgb(cor)
    # garante Times também para East Asia (Word às vezes troca)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)
    return r


def paragrafo(doc, texto="", **kw_fmt):
    p = _fmt(doc.add_paragraph(), **kw_fmt)
    if texto:
        _run(p, texto)
    return p


def titulo_secao(doc, texto):
    p = _fmt(doc.add_paragraph(), antes=18, depois=12, alinh=WD_ALIGN_PARAGRAPH.LEFT)
    _run(p, texto, negrito=True, cor=AZUL_ESCURO)
    return p


def sombrear(cell, hexa):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexa)
    cell._tc.get_or_add_tcPr().append(shd)


def bordas_finas(tabela):
    tbl_pr = tabela._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "1")  # 1/8 pt = 0,125 pt
        el.set(qn("w:color"), BORDA)
        borders.append(el)
    tbl_pr.append(borders)


def celula(cell, linhas, *, primeiro_negrito=False, cor="000000", fundo=None,
           alinh=WD_ALIGN_PARAGRAPH.LEFT):
    """Preenche a célula com uma lista de (texto, negrito) ou strings.
    Parágrafos compactos (sem o espaço-depois do corpo)."""
    if fundo:
        sombrear(cell, fundo)
    cell.paragraphs[0].text = ""
    for i, item in enumerate(linhas):
        texto, neg = item if isinstance(item, tuple) else (item, False)
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        _fmt(p, depois=2, alinh=alinh)
        if primeiro_negrito and i == 0:
            neg = True
        _run(p, texto, negrito=neg, cor=cor)


def celula_rica(cell, partes, *, fundo=None):
    """Célula com um parágrafo por item; item = lista de (texto, negrito)."""
    if fundo:
        sombrear(cell, fundo)
    cell.paragraphs[0].text = ""
    for i, spans in enumerate(partes):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        _fmt(p, depois=2, alinh=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for texto, neg in spans:
            _run(p, texto, negrito=neg)


def marcador(doc, texto):
    p = doc.add_paragraph(style="List Paragraph")
    pf = _fmt(p, depois=6, alinh=WD_ALIGN_PARAGRAPH.JUSTIFY).paragraph_format
    pf.left_indent = Pt(36)
    pf.first_line_indent = Pt(-18)
    _run(p, f"•\t{texto}")
    return p


def montar(dados: dict, template: Path, saida: Path):
    doc = Document(str(template))

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.left_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = TAM

    # limpa o corpo vazio do template (mantém cabeçalho/rodapé da seção)
    for p in list(doc.paragraphs):
        if not p.text.strip():
            p._element.getparent().remove(p._element)

    # ---- Bloco de capa (centralizado) ----
    p = _fmt(doc.add_paragraph(), depois=4, alinh=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, dados.get("titulo", "RELATÓRIO FINAL SIMPLIFICADO"), negrito=True, cor=AZUL_ESCURO)
    if dados.get("subtitulo"):
        p = _fmt(doc.add_paragraph(), depois=4, alinh=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, dados["subtitulo"], italico=True, cor=CINZA_CAPA1)
    if dados.get("unidade"):
        p = _fmt(doc.add_paragraph(), depois=4, alinh=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, dados["unidade"], italico=True)
    if dados.get("data"):
        p = _fmt(doc.add_paragraph(), depois=4, alinh=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, dados["data"], cor=CINZA_CAPA2)

    # ---- 1. Identificação ----
    titulo_secao(doc, "1. Identificação da Fiscalização")
    ident = dados.get("identificacao", [])
    if ident:
        t = doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        bordas_finas(t)
        for i, (rotulo, valor) in enumerate(ident):
            row = t.add_row()
            fundo = ZEBRA_AZUL if i % 2 == 0 else ZEBRA_CINZA
            celula(row.cells[0], [(rotulo, True)], fundo=fundo)
            celula(row.cells[1], [valor], fundo=fundo,
                   alinh=WD_ALIGN_PARAGRAPH.JUSTIFY)
        t.columns[0].width = Cm(5.4)
        t.columns[1].width = Cm(11.1)
        for row in t.rows:
            row.cells[0].width = Cm(5.4)
            row.cells[1].width = Cm(11.1)

    # ---- 2. Síntese ----
    titulo_secao(doc, "2. Síntese da Ação Fiscal")
    for par in dados.get("sintese", []):
        paragrafo(doc, par)

    # ---- 3. Notificações Lavradas ----
    titulo_secao(doc, "3. Notificações Lavradas")
    notifs = dados.get("notificacoes", [])
    if notifs:
        t = doc.add_table(rows=1, cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        bordas_finas(t)
        cab = t.rows[0].cells
        for c, txt in zip(cab, ("Notificação", "Itens notificados (resumo)", "Lavratura")):
            celula(c, [(txt, True)], cor=BRANCO, fundo=AZUL_ESCURO,
                   alinh=WD_ALIGN_PARAGRAPH.CENTER)
        for i, n in enumerate(notifs):
            row = t.add_row()
            fundo = ZEBRA_AZUL if i % 2 == 0 else ZEBRA_CINZA
            cod = [(n["codigo"], True)]
            if n.get("tipo"):
                cod.append((n["tipo"], False))
            celula(row.cells[0], cod, fundo=fundo)
            celula(row.cells[1], [n.get("itens", "")], fundo=fundo,
                   alinh=WD_ALIGN_PARAGRAPH.JUSTIFY)
            celula(row.cells[2], [n.get("lavrada", "")], fundo=fundo,
                   alinh=WD_ALIGN_PARAGRAPH.CENTER)
        larguras = (Cm(4.0), Cm(9.7), Cm(2.8))
        for j, w in enumerate(larguras):
            t.columns[j].width = w
        for row in t.rows:
            for j, w in enumerate(larguras):
                row.cells[j].width = w

    # ---- 4. Autos de Infração Lavrados ----
    titulo_secao(doc, "4. Autos de Infração Lavrados")
    temas = dados.get("autos_temas", [])
    if not temas:
        paragrafo(doc, dados.get(
            "sem_autos",
            "Não há autos de infração transmitidos no Sistema Auditor até a data deste relatório."))
    else:
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        bordas_finas(t)
        cab = t.rows[0].cells
        for c, txt in zip(cab, ("Auto de Infração", "Irregularidade constatada")):
            celula(c, [(txt, True)], cor=BRANCO, fundo=AZUL_ESCURO,
                   alinh=WD_ALIGN_PARAGRAPH.CENTER)
        zebra = 0
        for tema in temas:
            row = t.add_row()
            row.cells[0].merge(row.cells[1])
            celula(row.cells[0], [(tema["tema"], True)], cor=BRANCO, fundo=AZUL_MEDIO)
            for auto in tema["autos"]:
                row = t.add_row()
                fundo = ZEBRA_AZUL if zebra % 2 == 0 else ZEBRA_CINZA
                zebra += 1
                celula(row.cells[0], [
                    (f"Nº {auto['numero']}", True),
                    (f"Ementa {auto['ementa']}", False),
                    (auto["fundamento"], False),
                ], fundo=fundo)
                celula_rica(row.cells[1], [
                    [(auto["descricao"], False)],
                    [("Constatação: ", True), (auto["constatacao"], False)],
                ], fundo=fundo)
        larguras = (Cm(4.6), Cm(11.9))
        for j, w in enumerate(larguras):
            t.columns[j].width = w
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                cell.width = larguras[min(j, 1)]
        if dados.get("autos_total"):
            p = _fmt(doc.add_paragraph(), antes=8, depois=10)
            _run(p, dados["autos_total"], negrito=True)

    # ---- 5. Interdições e Embargos ----
    titulo_secao(doc, "5. Interdições e Embargos")
    inter = dados.get("interdicoes", [])
    if not inter:
        inter = ["Não houve lavratura de termo de interdição ou embargo nesta ação fiscal."]
    for par in inter:
        paragrafo(doc, par)

    # ---- 6. Observações/Pendências ----
    titulo_secao(doc, "6. Observações/Pendências")
    obs = dados.get("observacoes", [])
    if not obs:
        paragrafo(doc, "Nenhuma pendência identificada.")
    for item in obs:
        marcador(doc, item)

    # ---- 7. Encerramento e assinatura ----
    titulo_secao(doc, "7. Auditores-Fiscais do Trabalho Envolvidos")
    paragrafo(doc, "É o relatório.")
    for nome, cargo in dados.get("auditores", []):
        p = _fmt(doc.add_paragraph(), antes=24, depois=4, alinh=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, "_" * 40)
        p = _fmt(doc.add_paragraph(), depois=4, alinh=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, nome, negrito=True)
        p = _fmt(doc.add_paragraph(), depois=4, alinh=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, cargo)

    doc.save(str(saida))
    print(f"OK docx: {saida}")


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        return 1
    entrada = Path(sys.argv[1])
    dados = json.loads(entrada.read_text(encoding="utf-8"))
    saida = Path(sys.argv[2]) if len(sys.argv) > 2 else entrada.with_suffix(".docx")
    template = Path(__file__).parent / "template-cabecalho.docx"
    if not template.exists():
        print(f"ERRO: template não encontrado: {template}")
        return 1
    montar(dados, template, saida)
    return 0


if __name__ == "__main__":
    sys.exit(main())
