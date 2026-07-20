"""modelo_docx.py — biblioteca do padrão visual de .docx do AFT Toolkit.

Todo documento .docx gerado pelo toolkit (fora os que têm template oficial
próprio, como o RT de interdição e a Relação de autos) usa este módulo. Ele
abre o template oficial com o cabeçalho da auditoria (AFT/SIT) e oferece as
peças do modelo "Relatório de Fiscalização / Indícios":

    import modelo_docx as m
    doc = m.novo_documento()
    m.capa(doc, "TÍTULO DO DOCUMENTO", subtitulo="descrição em itálico",
           unidade="EMPRESA LTDA — CNPJ 00.000.000/0000-00",
           data="Goiânia-GO, 20 de julho de 2026")
    m.titulo_secao(doc, "1. Primeira Seção")
    m.paragrafo(doc, "Corpo justificado, Times 12, entrelinhas 1,15.")
    m.subtitulo(doc, "1.1 Subtítulo de nível 2")
    m.marcador(doc, "item de lista com marcador")
    m.tabela_rotulo_valor(doc, [("Empresa", "..."), ("CNPJ", "...")])
    t = m.nova_tabela(doc, ["Coluna A", "Coluna B"], larguras_cm=(5, 11.5))
    m.linha_subcabecalho(t, "Subgrupo dentro da tabela")
    m.linha_dados(t, ["valor A", "valor B"])          # zebra automática
    m.assinatura(doc, "Nome do Auditor", "Auditor-Fiscal do Trabalho — CIF 000000")
    doc.save("saida.docx")

Especificação: Times New Roman 12 em tudo; página A4 com margens sup/inf/dir
2 cm e esq 2,5 cm; paleta azul #1F3864 (títulos, cabeçalho de tabela) e
#2E5496 (subtítulos, subcabeçalhos); zebra #EBF3FB/#F5F5F5; bordas #AAAAAA
finas; entrelinhas 1,15; espaço-depois: título 12pt, subtítulo 6pt, corpo
10pt, lista 6pt, capa 4pt; espaço-antes: título 18pt, subtítulo 6pt.
O cabeçalho do template NUNCA é alterado.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AZUL_ESCURO = "1F3864"   # título principal, títulos de seção, cabeçalho de tabela
AZUL_MEDIO = "2E5496"    # subtítulos nível 2, subcabeçalhos de tabela
BRANCO = "FFFFFF"
PRETO = "000000"
CINZA_CAPA1 = "444444"   # subtítulo da capa
CINZA_CAPA2 = "555555"   # data da capa
ZEBRA_AZUL = "EBF3FB"
ZEBRA_CINZA = "F5F5F5"
BORDA = "AAAAAA"

FONTE = "Times New Roman"
TAM = Pt(12)

TEMPLATE = Path(__file__).parent / "template-cabecalho.docx"

ESQUERDA = WD_ALIGN_PARAGRAPH.LEFT
CENTRO = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFICADO = WD_ALIGN_PARAGRAPH.JUSTIFY


def rgb(hexa: str) -> RGBColor:
    return RGBColor.from_string(hexa)


# ---------------------------------------------------------------- fundação
def novo_documento(template=None) -> Document:
    """Documento novo sobre o template oficial (cabeçalho AFT/SIT preservado),
    margens e estilo Normal (Times 12) já aplicados."""
    tpl = Path(template) if template else TEMPLATE
    if not tpl.exists():
        raise FileNotFoundError(f"template do modelo-docx não encontrado: {tpl}")
    doc = Document(str(tpl))
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = TAM
    for p in list(doc.paragraphs):          # corpo do template vem vazio
        if not p.text.strip():
            p._element.getparent().remove(p._element)
    return doc


def fmt(p, *, antes=0, depois=10, alinh=JUSTIFICADO):
    """Aplica o espaçamento padrão (entrelinhas 1,15) a um parágrafo."""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(antes)
    pf.space_after = Pt(depois)
    pf.alignment = alinh
    return p


def run(p, texto, *, negrito=False, italico=False, cor=PRETO):
    """Acrescenta um run Times 12 formatado ao parágrafo."""
    r = p.add_run(texto)
    r.font.name = FONTE
    r.font.size = TAM
    r.bold = negrito
    r.italic = italico
    r.font.color.rgb = rgb(cor)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)
    return r


# ---------------------------------------------------------------- blocos
def capa(doc, titulo, subtitulo=None, unidade=None, data=None):
    """Bloco de capa centralizado: título navy negrito, subtítulo itálico
    cinza, linha da unidade fiscalizada em itálico, data em cinza."""
    p = fmt(doc.add_paragraph(), depois=4, alinh=CENTRO)
    run(p, titulo, negrito=True, cor=AZUL_ESCURO)
    if subtitulo:
        p = fmt(doc.add_paragraph(), depois=4, alinh=CENTRO)
        run(p, subtitulo, italico=True, cor=CINZA_CAPA1)
    if unidade:
        p = fmt(doc.add_paragraph(), depois=4, alinh=CENTRO)
        run(p, unidade, italico=True)
    if data:
        p = fmt(doc.add_paragraph(), depois=4, alinh=CENTRO)
        run(p, data, cor=CINZA_CAPA2)


def titulo_secao(doc, texto):
    """Título de seção ("1. ...", "2. ..."): negrito navy, 18pt antes/12 depois."""
    p = fmt(doc.add_paragraph(), antes=18, depois=12, alinh=ESQUERDA)
    run(p, texto, negrito=True, cor=AZUL_ESCURO)
    return p


def subtitulo(doc, texto):
    """Subtítulo de nível 2 ("2.1 ..."): negrito azul médio, 6pt antes/depois."""
    p = fmt(doc.add_paragraph(), antes=6, depois=6, alinh=ESQUERDA)
    run(p, texto, negrito=True, cor=AZUL_MEDIO)
    return p


def paragrafo(doc, texto="", *, antes=0, depois=10, alinh=JUSTIFICADO,
              negrito=False, italico=False, cor=PRETO):
    p = fmt(doc.add_paragraph(), antes=antes, depois=depois, alinh=alinh)
    if texto:
        run(p, texto, negrito=negrito, italico=italico, cor=cor)
    return p


def marcador(doc, texto):
    """Item de lista com marcador • (recuo 36pt, deslocado −18pt)."""
    p = doc.add_paragraph(style="List Paragraph")
    pf = fmt(p, depois=6).paragraph_format
    pf.left_indent = Pt(36)
    pf.first_line_indent = Pt(-18)
    run(p, f"•\t{texto}")
    return p


def assinatura(doc, nome, cargo, *, fecho="É o relatório."):
    """Fecho + linha de assinatura + nome/cargo, centralizados."""
    if fecho:
        paragrafo(doc, fecho)
    p = fmt(doc.add_paragraph(), antes=24, depois=4, alinh=CENTRO)
    run(p, "_" * 40)
    p = fmt(doc.add_paragraph(), depois=4, alinh=CENTRO)
    run(p, nome, negrito=True)
    p = fmt(doc.add_paragraph(), depois=4, alinh=CENTRO)
    run(p, cargo)


# ---------------------------------------------------------------- tabelas
def sombrear(cell, hexa):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexa)
    cell._tc.get_or_add_tcPr().append(shd)


def bordas_finas(tabela):
    tbl_pr = tabela._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "1")            # 1/8 pt = 0,125 pt
        el.set(qn("w:color"), BORDA)
        borders.append(el)
    tbl_pr.append(borders)


def celula(cell, linhas, *, primeiro_negrito=False, cor=PRETO, fundo=None,
           alinh=ESQUERDA):
    """Preenche a célula com parágrafos compactos. `linhas` = lista de strings
    ou de tuplas (texto, negrito)."""
    if fundo:
        sombrear(cell, fundo)
    cell.paragraphs[0].text = ""
    for i, item in enumerate(linhas):
        texto, neg = item if isinstance(item, tuple) else (item, False)
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        fmt(p, depois=2, alinh=alinh)
        if primeiro_negrito and i == 0:
            neg = True
        run(p, texto, negrito=neg, cor=cor)


def celula_rica(cell, partes, *, fundo=None):
    """Célula com um parágrafo por item; item = lista de (texto, negrito) —
    permite misturar negrito e normal na mesma linha (ex.: rótulo + valor)."""
    if fundo:
        sombrear(cell, fundo)
    cell.paragraphs[0].text = ""
    for i, spans in enumerate(partes):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        fmt(p, depois=2, alinh=JUSTIFICADO)
        for texto, neg in spans:
            run(p, texto, negrito=neg)


def _larguras(tabela, larguras_cm):
    if not larguras_cm:
        return
    ws = [Cm(w) for w in larguras_cm]
    for j, w in enumerate(ws):
        tabela.columns[j].width = w
    for row in tabela.rows:
        for j, cell in enumerate(row.cells):
            cell.width = ws[min(j, len(ws) - 1)]


def tabela_rotulo_valor(doc, linhas, larguras_cm=(5.4, 11.1)):
    """Tabela de identificação: 1ª coluna rótulo em negrito, 2ª valor;
    zebra e bordas finas, sem linha de cabeçalho."""
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    bordas_finas(t)
    for i, (rotulo, valor) in enumerate(linhas):
        row = t.add_row()
        fundo = ZEBRA_AZUL if i % 2 == 0 else ZEBRA_CINZA
        celula(row.cells[0], [(rotulo, True)], fundo=fundo)
        celula(row.cells[1], [valor], fundo=fundo, alinh=JUSTIFICADO)
    _larguras(t, larguras_cm)
    return t


def nova_tabela(doc, cabecalho, larguras_cm=None):
    """Tabela de dados com linha de cabeçalho navy (texto branco negrito,
    centralizado) e bordas finas. Preencher com linha_subcabecalho() e
    linha_dados(); a zebra é automática (reiniciável por subcabeçalho não é —
    ela alterna pelas linhas de dados adicionadas)."""
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    bordas_finas(t)
    for c, txt in zip(t.rows[0].cells, cabecalho):
        celula(c, [(txt, True)], cor=BRANCO, fundo=AZUL_ESCURO, alinh=CENTRO)
    t._modelo_larguras = larguras_cm
    t._modelo_zebra = 0
    if larguras_cm:
        _larguras(t, larguras_cm)
    return t


def linha_subcabecalho(tabela, texto):
    """Linha de subgrupo dentro da tabela: fundo azul médio, branco negrito,
    mesclada em toda a largura."""
    row = tabela.add_row()
    cell = row.cells[0]
    for outra in row.cells[1:]:
        cell = cell.merge(outra)
    celula(cell, [(texto, True)], cor=BRANCO, fundo=AZUL_MEDIO)
    if getattr(tabela, "_modelo_larguras", None):
        _larguras(tabela, tabela._modelo_larguras)
    return row


def linha_dados(tabela, celulas):
    """Linha de dados com zebra automática. Cada item de `celulas` pode ser:
    string · lista de strings/tuplas (vira celula()) · dict {"rica": [...]}
    (vira celula_rica())."""
    row = tabela.add_row()
    fundo = ZEBRA_AZUL if tabela._modelo_zebra % 2 == 0 else ZEBRA_CINZA
    tabela._modelo_zebra += 1
    for cell, item in zip(row.cells, celulas):
        if isinstance(item, dict) and "rica" in item:
            celula_rica(cell, item["rica"], fundo=fundo)
        elif isinstance(item, list):
            celula(cell, item, fundo=fundo)
        else:
            celula(cell, [item], fundo=fundo, alinh=JUSTIFICADO)
    if getattr(tabela, "_modelo_larguras", None):
        _larguras(tabela, tabela._modelo_larguras)
    return row
