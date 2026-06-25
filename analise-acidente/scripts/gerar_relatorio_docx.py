# -*- coding: utf-8 -*-
"""
Monta o Relatorio de Analise de Acidente (.docx) a partir de um arquivo JSON de conteudo.

Uso:
    python gerar_relatorio_docx.py conteudo.json

O JSON descreve o relatorio; este script cuida da formatacao (margens, fonte, cores,
cabecalho, tabela de identificacao, secoes, subtitulos, bullets e blocos de fator SFIT).
A skill /analise-acidente preenche o conteudo; a formatacao fica padronizada aqui.

Esquema do JSON:
{
  "saida": "caminho\\\\Relatorio.docx",          # obrigatorio
  "titulo": "RELATORIO DE ANALISE DE ACIDENTE DO TRABALHO",
  "subtitulo": "Acidente do trabalho tipico com obito",   # opcional
  "identificacao": [["Empregador","..."], ["CNPJ","..."]],# linhas da tabela de capa
  "secoes": [
     {"titulo": "1. DESCRICAO DO LOCAL DO ACIDENTE",
      "blocos": [
         {"t":"p",   "x":"paragrafo, aceita **negrito** inline"},
         {"t":"sub", "x":"Subtitulo em negrito"},
         {"t":"b",   "x":"item de lista (bullet)"},
         {"t":"fator","codigo":"251008","nome":"...","classe":"determinante","desc":"..."}
      ]}
  ],
  "rodape": "texto final em italico pequeno"             # opcional
}
Tipos de bloco: "p" paragrafo | "sub" subtitulo | "b" bullet | "fator" fator SFIT.
"""
import sys, os, json, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AZUL = RGBColor(0x1F, 0x3A, 0x5F)

def add_runs(p, texto, size=11, base_bold=False):
    """Interpreta **negrito** inline e adiciona os runs ao paragrafo."""
    for parte in re.split(r'(\*\*[^*]+\*\*)', texto):
        if parte.startswith('**') and parte.endswith('**'):
            r = p.add_run(parte[2:-2]); r.bold = True
        else:
            r = p.add_run(parte); r.bold = base_bold
        r.font.size = Pt(size)

def main(json_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        d = json.load(f)

    doc = Document()
    nrm = doc.styles['Normal']; nrm.font.name = 'Calibri'; nrm.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3.0); s.right_margin = Cm(2.5)

    def par(texto, bold=False, italic=False, align='just', size=11, after=6, before=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = {'just': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                       'left': WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
        if italic:
            r = p.add_run(texto); r.italic = True; r.font.size = Pt(size); r.bold = bold
        else:
            add_runs(p, texto, size=size, base_bold=bold)
        return p

    # ---- cabecalho ----
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run('MINISTÉRIO DO TRABALHO E EMPREGO'); r.bold = True; r.font.size = Pt(12)
    c2 = doc.add_paragraph(); c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c2.add_run('Inspeção do Trabalho - Auditoria-Fiscal do Trabalho'); r.font.size = Pt(10)
    doc.add_paragraph()
    tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tt.add_run(d.get('titulo', 'RELATORIO DE ANALISE DE ACIDENTE DO TRABALHO'))
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = AZUL
    if d.get('subtitulo'):
        st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = st.add_run(d['subtitulo']); r.italic = True; r.font.size = Pt(11)
    doc.add_paragraph()

    # ---- tabela de identificacao ----
    ident = d.get('identificacao') or []
    if ident:
        tb = doc.add_table(rows=0, cols=2); tb.style = 'Light Grid Accent 1'
        for k, v in ident:
            cells = tb.add_row().cells
            cells[0].paragraphs[0].add_run(str(k)).bold = True
            cells[1].paragraphs[0].add_run(str(v))
            for cc in cells:
                for pp in cc.paragraphs: pp.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    # ---- secoes ----
    for sec in d.get('secoes', []):
        h = doc.add_heading(level=1)
        rr = h.add_run(sec.get('titulo', '')); rr.font.color.rgb = AZUL; rr.font.size = Pt(13)
        for bloco in sec.get('blocos', []):
            t = bloco.get('t', 'p')
            if t == 'p':
                par(bloco.get('x', ''))
            elif t == 'sub':
                par(bloco.get('x', ''), bold=True, align='left', after=2, before=6)
            elif t == 'b':
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                add_runs(p, bloco.get('x', ''))
            elif t == 'fator':
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(4); p.paragraph_format.line_spacing = 1.15
                rr = p.add_run('Fator Causal %s - %s' % (bloco.get('codigo', ''), bloco.get('nome', '')))
                rr.bold = True; rr.font.size = Pt(11)
                p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p2.paragraph_format.space_after = Pt(6); p2.paragraph_format.line_spacing = 1.15
                lab = p2.add_run('Classificacao (%s). Descricao: ' % bloco.get('classe', 'a confirmar'))
                lab.bold = True; lab.font.size = Pt(11)
                add_runs(p2, bloco.get('desc', ''))

    if d.get('rodape'):
        doc.add_paragraph()
        par(d['rodape'], italic=True, size=9, after=0)

    saida = d['saida']
    os.makedirs(os.path.dirname(saida), exist_ok=True) if os.path.dirname(saida) else None
    doc.save(saida)
    print('SALVO:', saida)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Uso: python gerar_relatorio_docx.py conteudo.json'); sys.exit(1)
    main(sys.argv[1])
