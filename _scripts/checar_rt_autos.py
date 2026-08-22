#!/usr/bin/env python3
"""
checar_rt_autos.py - coerencia entre as irregularidades do RT (.docx) e os autos
(autos.md). Cobre o RT nos dois formatos do montar_rt.py: "topico" (secao 4
unica) e "objeto" (um bloco de irregularidades por objeto interditado).

Ao editar o RT (trocar ementas por itens da NR, tirar/incluir irregularidades), o RT e
o autos.md podem desalinhar. Este verificador compara os dois e SINALIZA divergencias
para o AFT revisar (nao bloqueia, nao corrige - so aponta).

Sinais comparados (robustos, de baixo falso-positivo):
  1. CONTAGEM: nº de irregularidades na Secao 4 do RT x nº de autos no autos.md.
  2. NRs ENVOLVIDAS: conjunto de NRs citadas na Secao 4 x nos autos (pega casos como
     "NR-35 no RT mas sem auto de NR-35", ou "NR-01 nos autos mas nao no RT").
  3. CODIGOS DE EMENTA: se o RT trouxer codigos (formato XXXXXX-X), compara os conjuntos.

NAO faz casamento por subitem (33.3.1 x 33.3.2 etc.) porque a relacao item-de-NR x
ementa e muitos-para-muitos e geraria ruido. Lista os dois lados para conferencia manual.

Uso:
    python checar_rt_autos.py "<RT.docx>" "<autos.md>"
Exit 0 = sem divergencia detectada; 1 = ha divergencia (revisar).
"""

try:  # ticket automatico de erro (ver _scripts/erro_ticket.py e a skill /aft-erro)
    import sys as _sys
    from pathlib import Path as _Path
    _aqui = _Path(__file__).resolve()
    for _p in (_aqui.parent, *(_a / "_scripts" for _a in _aqui.parents)):
        if (_p / "erro_ticket.py").is_file():
            _sys.path.insert(0, str(_p))
            from erro_ticket import ativar as _ativar_ticket
            _ativar_ticket(__file__)
            break
except Exception:
    pass

import os
import re
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

CODE_RE = re.compile(r"\d{6}-\d")
NR_RE = re.compile(r"NR[-\s]?0*(\d{1,2})", re.IGNORECASE)

# A NR-3 e a norma da PROPRIA medida (metodologia de caracterizacao do grave e
# iminente risco). Ela aparece no texto fixo da Secao 4 de todo RT, mas nunca e
# a ementa de um auto - por isso fica fora da comparacao de NRs.
NRS_IGNORADAS = {3}


def nrs(texto):
    return {int(n) for n in NR_RE.findall(texto)} - NRS_IGNORADAS


def codes(texto):
    return set(CODE_RE.findall(texto))


RE_LEGENDA = re.compile(r"^\s*(fotografia|foto|figura|imagem)\b", re.IGNORECASE)


def eh_legenda(texto):
    """Legenda de foto inserida no corpo do RT nao e irregularidade."""
    return bool(RE_LEGENDA.match(texto))


RE_IRREG_PDF = re.compile(r"IRREGULARIDADE(?:S)?\s*:?", re.IGNORECASE)
RE_FATOR_PDF = re.compile(r"FATOR(?:ES)?\s+DE\s+RISCO", re.IGNORECASE)
RE_RODAPE_PDF = re.compile(r"^T\.\s*(Interdi|Embargo)", re.IGNORECASE)


def ler_secao4_rt_pdf(rt_path):
    """Mesma leitura, para RT em PDF.

    No Modo B da skill (Termo ja lavrado, so faltam os autos) o RT quase nunca e
    o .docx que o montar_rt.py gerou: e o PDF impresso pelo Sistema Auditor, que
    e o que o AFT tem em maos. Antes desta funcao o script chamava
    docx.Document() direto no PDF e morria com PackageNotFoundError -- traceback
    cru, ticket automatico, e a conferencia RT x autos simplesmente nao rodava
    justamente no caso mais comum.

    Sem camada de texto (PDF escaneado), devolve (None, None): quem chama ja
    trata isso como "nao encontrei o bloco de irregularidades".
    """
    try:
        import pdfplumber
    except ImportError:
        print("AVISO: RT em PDF exige a biblioteca pdfplumber "
              "(pip install pdfplumber). Conferencia nao realizada.",
              file=sys.stderr)
        return None, None

    with pdfplumber.open(rt_path) as pdf:
        linhas = []
        for pagina in pdf.pages:
            for linha in (pagina.extract_text() or "").splitlines():
                linha = linha.strip()
                # o rodape se repete em toda pagina e traz o numero do termo,
                # que o CODE_RE confundiria com codigo de ementa
                if linha and not RE_RODAPE_PDF.match(linha):
                    linhas.append(linha)

    inicios = [i for i, t in enumerate(linhas) if RE_IRREG_PDF.fullmatch(t)]
    if not inicios:
        return None, None

    itens, textos = [], []
    for ini in inicios:
        buffer = ""
        for t in linhas[ini + 1:]:
            if RE_FATOR_PDF.search(t):
                break
            textos.append(t)
            # no PDF a ementa quebra em varias linhas; um item novo comeca
            # sempre por um codigo XXXXXX-X no inicio da linha
            if CODE_RE.match(t):
                if buffer:
                    itens.append(buffer.strip())
                buffer = t
            elif buffer:
                buffer += " " + t
        if buffer:
            itens.append(buffer.strip())

    itens = [t for t in itens if t and not eh_legenda(t)]
    # ementa repetida entre objetos rende UM auto so (regra 7.1 da skill)
    vistos, unicos = set(), []
    for t in itens:
        m = CODE_RE.search(t)
        chave = m.group(0) if m else t
        if chave not in vistos:
            vistos.add(chave)
            unicos.append(t)
    return unicos, "\n".join(textos)


def ler_secao4_rt(rt_path):
    """Retorna (itens_de_irregularidade, texto_dos_blocos_de_irregularidade).

    Aceita RT em .docx e em .pdf (ver ler_secao4_rt_pdf).

    Nos .docx, aceita os tres layouts de RT:
      - topico atual: os titulos usam numeracao AUTOMATICA do Word, entao o
        texto e so "IRREGULARIDADE(S):" (sem o "4.") e as irregularidades sao
        paragrafos de prosa ANTES do bloco fixo da metodologia da NR-3;
      - topico anterior: titulos com "4." digitado e ementas em lista (numPr)
        DEPOIS do bloco da metodologia;
      - POR OBJETO (formato "objeto" do montar_rt.py): nao ha secao 4; cada
        objeto tem seu rotulo "IRREGULARIDADE(S):" seguido das ementas, ate o
        rotulo "FATOR(ES) DE RISCO...". Detectado por haver MAIS DE UM rotulo
        de irregularidade; as ementas de todos os objetos sao somadas.
    Em todos, as tabelas 3.1/3.2/3.3 ficam de fora: doc.paragraphs nao inclui o
    conteudo de tabelas.
    """
    if os.path.splitext(rt_path)[1].lower() == ".pdf":
        return ler_secao4_rt_pdf(rt_path)

    import docx
    try:
        doc = docx.Document(rt_path)
    except Exception as e:
        # arquivo que nao e um .docx valido (RT salvo com outro formato, arquivo
        # truncado): mensagem util em vez de traceback do python-docx
        print(f"ERRO: nao consegui abrir o RT como .docx: {e}", file=sys.stderr)
        print("Formatos aceitos: .docx (gerado pelo montar_rt.py) e .pdf "
              "(impresso pelo Sistema Auditor).", file=sys.stderr)
        return None, None
    paras = doc.paragraphs

    def eh_titulo_irregularidade(t):
        return "IRREGULARIDADE" in t and len(t) < 40

    def eh_titulo_fator(t):
        return ("FATOR" in t or "RISCO" in t) and "RELACIONADO" in t

    titulos = [i for i, p in enumerate(paras)
               if eh_titulo_irregularidade(p.text.strip().upper())]
    if not titulos:
        return None, None

    if len(titulos) > 1:                       # RT por objeto
        itens, textos = [], []
        for ini in titulos:
            for j in range(ini + 1, len(paras)):
                t = paras[j].text.strip()
                if eh_titulo_fator(t.upper()):
                    break
                textos.append(paras[j].text)
                if t and not eh_legenda(t):
                    itens.append(t)
        # ementa que atinge mais de um objeto se repete no bloco de cada um,
        # mas rende UM auto so (regra 7.1 da skill) - deduplica para a contagem
        vistos, unicos = set(), []
        for t in itens:
            m = CODE_RE.search(t)
            chave = m.group(0) if m else t
            if chave not in vistos:
                vistos.add(chave)
                unicos.append(t)
        return unicos, "\n".join(textos)

    ini = titulos[0]
    fim = metodologia = None
    for i in range(ini + 1, len(paras)):
        t = paras[i].text.strip().upper()
        if not t:
            continue
        if metodologia is None and t.startswith("METODOLOGIA PRESCRITA"):
            metodologia = i
        if eh_titulo_fator(t):
            fim = i
            break

    sec = paras[ini + 1:(fim if fim is not None else len(paras))]
    texto = "\n".join(p.text for p in sec)

    # formato atual: o que vem entre o titulo e a metodologia
    if metodologia is not None:
        antes = [p.text.strip() for p in paras[ini + 1:metodologia]
                 if p.text.strip() and not eh_legenda(p.text)]
        if antes:
            return antes, texto

    # formato anterior: paragrafos de lista em qualquer ponto da secao
    itens = [p.text.strip() for p in sec
             if p.text.strip() and p._p.pPr is not None and p._p.pPr.numPr is not None]
    return itens, texto


def ler_autos(autos_path):
    """Retorna lista de (codigo, descricao) e o texto completo."""
    texto = open(autos_path, encoding="utf-8").read()
    blocos = re.split(r"=== AUTO(?: DE INFRAÇÃO)? #\d+ ===", texto)
    autos = []
    for b in blocos:
        m = re.search(r"Ementa:\s*(\d{6}-\d)\s*-\s*(.*)", b)
        if m:
            autos.append((m.group(1), m.group(2).strip()))
    return autos, texto


def main():
    if len(sys.argv) != 3:
        print("uso: python checar_rt_autos.py <RT.docx|RT.pdf> <autos.md>", file=sys.stderr)
        sys.exit(2)
    rt_path, autos_path = sys.argv[1], sys.argv[2]
    for p in (rt_path, autos_path):
        if not os.path.isfile(p):
            print(f"ERRO: arquivo nao encontrado: {p}", file=sys.stderr)
            sys.exit(2)

    itens_rt, texto_rt = ler_secao4_rt(rt_path)
    if itens_rt is None:
        print("ERRO: nao encontrei bloco de 'IRREGULARIDADE(S)' no RT.", file=sys.stderr)
        sys.exit(2)
    autos, texto_autos = ler_autos(autos_path)

    nrs_rt, nrs_au = nrs(texto_rt), nrs(texto_autos)
    cod_rt, cod_au = codes(texto_rt), codes(texto_autos)

    div = []  # divergencias

    print("=" * 64)
    print("  COERENCIA RT  x  autos.md")
    print("=" * 64)
    print(f"Irregularidades no RT : {len(itens_rt)}")
    print(f"autos.md              : {len(autos)} auto(s)")
    if len(itens_rt) != len(autos):
        div.append(f"Contagem divergente: RT={len(itens_rt)} x autos={len(autos)}.")

    print("\nNRs no RT     : " + (", ".join(f"NR-{n:02d}" for n in sorted(nrs_rt)) or "-"))
    print("NRs nos autos : " + (", ".join(f"NR-{n:02d}" for n in sorted(nrs_au)) or "-"))
    so_rt = nrs_rt - nrs_au
    so_au = nrs_au - nrs_rt
    if so_rt:
        div.append("NR(s) no RT sem auto correspondente: "
                   + ", ".join(f"NR-{n:02d}" for n in sorted(so_rt)))
    if so_au:
        div.append("NR(s) nos autos sem item no RT: "
                   + ", ".join(f"NR-{n:02d}" for n in sorted(so_au)))

    if cod_rt:
        print("\nO RT contem codigos de ementa - comparando codigos:")
        c_so_au = cod_au - cod_rt
        c_so_rt = cod_rt - cod_au
        if c_so_rt:
            div.append("Ementa(s) no RT sem auto: " + ", ".join(sorted(c_so_rt)))
        if c_so_au:
            div.append("Ementa(s) nos autos sem respaldo no RT: " + ", ".join(sorted(c_so_au)))
    else:
        print("\n(O RT usa itens de NR, sem codigos de ementa - comparacao por NR + contagem.)")

    print("\n--- Autos (autos.md) ---")
    for i, (c, d) in enumerate(autos, 1):
        print(f"  {i:>2}. {c}  {d[:70]}")
    print("\n--- Irregularidades (RT) ---")
    for i, t in enumerate(itens_rt, 1):
        print(f"  {i:>2}. {t[:80]}")

    print("\n" + "-" * 64)
    if div:
        print(f"DIVERGENCIAS ({len(div)}) - revise:")
        for d in div:
            print("  ! " + d)
        print("\nObs.: contagem e NR sao sinais; confira a lista acima item a item.")
        sys.exit(1)
    print("Sem divergencias de contagem/NR/codigo. Confira mesmo assim a lista acima.")
    sys.exit(0)


if __name__ == "__main__":
    main()
